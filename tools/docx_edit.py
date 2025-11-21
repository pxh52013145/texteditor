import argparse
from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from docx.shared import RGBColor

def replace_in_runs(runs, old, new):
    for r in runs:
        if r.text:
            r.text = r.text.replace(old, new)

def replace_text(doc, replacements):
    for p in doc.paragraphs:
        for old, new in replacements.items():
            replace_in_runs(p.runs, old, new)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for old, new in replacements.items():
                        replace_in_runs(p.runs, old, new)

def set_font(doc, name=None, size_pt=None, color_hex=None):
    def apply_run(r):
        f = r.font
        if name:
            f.name = name
        if size_pt:
            f.size = Pt(size_pt)
        if color_hex:
            f.color.rgb = RGBColor.from_string(color_hex)
    for p in doc.paragraphs:
        for r in p.runs:
            apply_run(r)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        apply_run(r)

def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("-i", "--input", required=True)
    ap.add_argument("-o", "--output", required=True)
    ap.add_argument("--replace", nargs="*", default=[])
    ap.add_argument("--font-name")
    ap.add_argument("--font-size", type=float)
    ap.add_argument("--font-color")
    ap.add_argument("--heading")
    ap.add_argument("--heading-level", type=int, default=1)
    ap.add_argument("--image")
    ap.add_argument("--image-width", type=float)
    ap.add_argument("--append-from")
    args = ap.parse_args()
    doc = Document(args.input)
    repl = {}
    for pair in args.replace:
        if "=" in pair:
            k, v = pair.split("=", 1)
            repl[k] = v
    if repl:
        replace_text(doc, repl)
    if args.font_name or args.font_size or args.font_color:
        set_font(doc, args.font_name, args.font_size, args.font_color)
    if args.heading:
        try:
            doc.add_heading(args.heading, level=args.heading_level)
        except Exception:
            p = doc.add_paragraph(args.heading)
            r = p.runs[0] if p.runs else p.add_run(args.heading)
            r.bold = True
            r.font.size = Pt(18)
    if args.image:
        if args.image_width:
            doc.add_picture(args.image, width=Inches(args.image_width))
        else:
            doc.add_picture(args.image)
    if args.append_from:
        with open(args.append_from, "r", encoding="utf-8") as fh:
            for line in fh:
                s = line.rstrip("\n")
                if not s:
                    doc.add_paragraph("")
                    continue
                if s.startswith("#"):
                    lvl = len(s) - len(s.lstrip('#'))
                    text = s.lstrip('#').strip()
                    try:
                        doc.add_heading(text, level=max(1, min(9, lvl)))
                    except Exception:
                        p = doc.add_paragraph(text)
                        r = p.runs[0] if p.runs else p.add_run(text)
                        r.bold = True
                        size_map = {1: 18, 2: 16, 3: 14, 4: 12}
                        r.font.size = Pt(size_map.get(lvl, 12))
                    continue
                if s.startswith("-"):
                    text = s.lstrip('-').strip()
                    try:
                        doc.add_paragraph(text, style='List Bullet')
                    except Exception:
                        doc.add_paragraph("• " + text)
                    continue
                doc.add_paragraph(s)
    doc.save(args.output)

if __name__ == "__main__":
    main()